"""Export compiled notes / gaps to DOCX and PDF."""

from __future__ import annotations

import io
import json
import os
import re
from typing import Any

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def _parse_gaps(note: dict) -> list[dict]:
    try:
        gaps = json.loads(note.get("gap_list_json") or "[]")
    except Exception:
        gaps = []
    return gaps if isinstance(gaps, list) else []


def _md_lines(markdown: str) -> list[str]:
    return (markdown or "").replace("\r\n", "\n").split("\n")


def _add_runs_with_bold(paragraph, text: str) -> None:
    # Simple **bold** support
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown: str, *, title: str = "Betula Notu") -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xC2, 0x65, 0x2A)

    in_code = False
    code_buf: list[str] = []

    for raw in _md_lines(markdown):
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_buf))
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, re.sub(r"^[-*]\s+", "", line))
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, re.sub(r"^\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def gaps_to_docx(gaps: list[dict], *, title: str = "Eksik Bilgiler") -> bytes:
    doc = Document()
    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xC2, 0x65, 0x2A)

    if not gaps:
        doc.add_paragraph("Henüz eksik bilgi kaydı yok.")
    else:
        for i, g in enumerate(gaps, 1):
            topic = (g.get("topic") or f"Konu {i}").strip()
            doc.add_heading(f"{i}. {topic}", level=1)
            reason = (g.get("reason") or "").strip()
            if reason:
                p = doc.add_paragraph()
                run = p.add_run("Neden: ")
                run.bold = True
                p.add_run(reason)
            summary = (g.get("summary") or "").strip()
            if summary:
                p = doc.add_paragraph()
                run = p.add_run("Özet: ")
                run.bold = True
                p.add_run(summary)
            sources = [s for s in (g.get("sources") or []) if isinstance(s, dict) and s.get("href")]
            if sources:
                doc.add_paragraph("Kaynaklar:").runs[0].bold = True
                for s in sources:
                    doc.add_paragraph(
                        f"{s.get('title') or s['href']} — {s['href']}",
                        style="List Bullet",
                    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_PDF_FONT_FILE: str | None = None
_PDF_FONT_CACHE: fitz.Font | None = None
_PDF_FONTNAME = "betula-tr"


def _resolve_pdf_fontfile() -> str | None:
    """Turkce destekli TTF bul (Windows / Linux)."""
    candidates = [
        # Windows
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\tahoma.ttf",
        # Linux common
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        # macOS
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def _pdf_font() -> fitz.Font:
    global _PDF_FONT_FILE, _PDF_FONT_CACHE
    if _PDF_FONT_CACHE is not None:
        return _PDF_FONT_CACHE
    _PDF_FONT_FILE = _resolve_pdf_fontfile()
    if _PDF_FONT_FILE:
        _PDF_FONT_CACHE = fitz.Font(fontfile=_PDF_FONT_FILE)
    else:
        _PDF_FONT_CACHE = fitz.Font("helv")
    return _PDF_FONT_CACHE


def _pdf_ensure_page_font(page: fitz.Page) -> str:
    """Sayfaya Unicode fontu gom; kullanilacak font adini dondur."""
    global _PDF_FONT_FILE
    if _PDF_FONT_FILE is None:
        _PDF_FONT_FILE = _resolve_pdf_fontfile()
    if _PDF_FONT_FILE:
        try:
            page.insert_font(fontname=_PDF_FONTNAME, fontfile=_PDF_FONT_FILE)
            return _PDF_FONTNAME
        except Exception:
            pass
    return "helv"


def _wrap_text_lines(text: str, *, width: float, fontsize: float) -> list[str]:
    """Kelime kaydirarak satir listesi uret."""
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    font = _pdf_font()
    out: list[str] = []

    for para in text.split("\n"):
        raw = para.strip()
        if not raw:
            out.append("")
            continue
        words = raw.split()
        if not words:
            out.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            if font.text_length(trial, fontsize=fontsize) <= width:
                current = trial
            else:
                out.append(current)
                current = word
                while font.text_length(current, fontsize=fontsize) > width and len(current) > 1:
                    cut = max(1, int(len(current) * width / max(font.text_length(current, fontsize=fontsize), 1)))
                    out.append(current[:cut])
                    current = current[cut:]
        out.append(current)
    return out


def _pdf_draw_lines(
    doc: fitz.Document,
    page: fitz.Page,
    lines: list[str],
    *,
    x: float,
    y: float,
    width: float,
    fontsize: float,
    margin: float,
    page_bottom: float,
    line_gap: float = 1.35,
    para_gap: float = 6.0,
) -> tuple[fitz.Page, float]:
    """Satirlari ust uste binmeden yazar; gerekirse yeni sayfa acar."""
    fontname = _pdf_ensure_page_font(page)
    line_h = fontsize * line_gap
    for line in lines:
        if line == "":
            y += para_gap * 0.6
            continue
        if y + line_h > page_bottom:
            page = doc.new_page(width=595, height=842)
            fontname = _pdf_ensure_page_font(page)
            y = margin
        page.insert_text(
            (x, y + fontsize * 0.85),
            line,
            fontsize=fontsize,
            fontname=fontname,
            color=(0.14, 0.1, 0.08),
        )
        y += line_h
    return page, y + para_gap


def markdown_to_pdf(markdown: str, *, title: str = "Betula Notu") -> bytes:
    # Font onbellegini hazirla
    _pdf_font()

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    margin = 50
    width = 595 - 2 * margin
    page_bottom = 842 - margin
    y = margin

    title_lines = _wrap_text_lines(title, width=width, fontsize=18)
    page, y = _pdf_draw_lines(
        doc, page, title_lines, x=margin, y=y, width=width, fontsize=18, margin=margin, page_bottom=page_bottom, para_gap=14
    )

    in_code = False
    for raw in _md_lines(markdown):
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            clean = line
            size = 9
        else:
            if not line.strip():
                y += 8
                continue
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            if re.match(r"^#\s+", clean):
                clean = re.sub(r"^#\s+", "", clean)
                size = 16
                y += 6
            elif re.match(r"^##\s+", clean):
                clean = re.sub(r"^##\s+", "", clean)
                size = 13
                y += 4
            elif re.match(r"^###\s+", clean):
                clean = re.sub(r"^###\s+", "", clean)
                size = 12
                y += 2
            elif re.match(r"^[-*]\s+", clean):
                clean = "• " + re.sub(r"^[-*]\s+", "", clean)
                size = 11
            elif re.match(r"^\d+\.\s+", clean):
                size = 11
            else:
                size = 11

        wrapped = _wrap_text_lines(clean, width=width, fontsize=size)
        page, y = _pdf_draw_lines(
            doc,
            page,
            wrapped,
            x=margin,
            y=y,
            width=width,
            fontsize=size,
            margin=margin,
            page_bottom=page_bottom,
            para_gap=5 if size >= 12 else 4,
        )

    out = io.BytesIO()
    doc.save(out)
    doc.close()
    return out.getvalue()


def gaps_to_pdf(gaps: list[dict], *, title: str = "Eksik Bilgiler") -> bytes:
    blocks: list[str] = []
    if not gaps:
        blocks.append("Henüz eksik bilgi kaydı yok.")
    else:
        for i, g in enumerate(gaps, 1):
            blocks.append(f"## {i}. {(g.get('topic') or 'Konu').strip()}")
            if g.get("reason"):
                blocks.append(f"**Neden:** {g['reason']}")
            if g.get("summary"):
                blocks.append(str(g["summary"]))
            sources = [s for s in (g.get("sources") or []) if isinstance(s, dict) and s.get("href")]
            if sources:
                blocks.append("Kaynaklar:")
                for s in sources:
                    blocks.append(f"- {s.get('title') or s['href']}: {s['href']}")
            blocks.append("")
    return markdown_to_pdf("\n\n".join(blocks), title=title)


def _safe_filename_base(title: str | None, fallback: str = "betula_note") -> str:
    """HTTP Content-Disposition icin ASCII-guvenli dosya adi uret."""
    raw = (title or fallback).strip() or fallback
    # Turkce/unicode karakterler header'da latin-1 hatasi verir
    ascii_only = raw.encode("ascii", "ignore").decode("ascii")
    cleaned = re.sub(r"[^\w\-]+", "_", ascii_only).strip("_")
    return (cleaned or fallback)[:60]


def build_export(
    note: dict,
    *,
    kind: str,
    fmt: str,
    doc_title: str | None = None,
) -> tuple[bytes, str, str]:
    """
    kind: note | gaps | full
    fmt: md | docx | pdf
    returns: (bytes, media_type, filename)
    """
    kind = (kind or "note").lower()
    fmt = (fmt or "docx").lower()
    gaps = _parse_gaps(note)
    md = note.get("markdown") or ""
    base = _safe_filename_base(doc_title, f"betula_{note.get('document_id') or 'note'}")

    if kind == "gaps":
        title = "Eksik Bilgiler — Betula"
        if fmt == "docx":
            return gaps_to_docx(gaps, title=title), (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ), f"{base}_eksik_bilgiler.docx"
        if fmt == "pdf":
            return gaps_to_pdf(gaps, title=title), "application/pdf", f"{base}_eksik_bilgiler.pdf"
        # md
        lines = ["# Eksik Bilgiler\n"]
        for i, g in enumerate(gaps, 1):
            lines.append(f"## {i}. {g.get('topic') or 'Konu'}\n")
            if g.get("reason"):
                lines.append(f"**Neden:** {g['reason']}\n")
            if g.get("summary"):
                lines.append(f"{g['summary']}\n")
        body = "\n".join(lines).encode("utf-8")
        return body, "text/markdown; charset=utf-8", f"{base}_eksik_bilgiler.md"

    if kind == "full":
        gap_md_parts = ["\n\n---\n\n# Eksik Bilgiler\n"]
        for i, g in enumerate(gaps, 1):
            gap_md_parts.append(f"\n## {i}. {g.get('topic') or 'Konu'}\n")
            if g.get("reason"):
                gap_md_parts.append(f"**Neden:** {g['reason']}\n")
            if g.get("summary"):
                gap_md_parts.append(f"{g['summary']}\n")
        md = md + "".join(gap_md_parts)

    title = "Master Sentez — Betula"
    if fmt == "docx":
        return markdown_to_docx(md, title=title), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ), f"{base}_master_sentez.docx"
    if fmt == "pdf":
        return markdown_to_pdf(md, title=title), "application/pdf", f"{base}_master_sentez.pdf"
    return md.encode("utf-8"), "text/markdown; charset=utf-8", f"{base}_master_sentez.md"
