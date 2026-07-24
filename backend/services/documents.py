"""Document text extraction (PyMuPDF + DOCX)."""

from __future__ import annotations

import hashlib
import io
import os
from typing import BinaryIO, Union

import fitz  # PyMuPDF
from docx import Document


FileLike = Union[BinaryIO, bytes]


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def checksum_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def get_pdf_text_from_bytes(data: bytes) -> str:
    text_parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def get_docx_text_from_bytes(data: bytes) -> str:
    text = ""
    doc = Document(io.BytesIO(data))
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text.strip()


def extract_document_text(filename: str, data: bytes) -> dict:
    extension = get_file_extension(filename)
    if extension == ".pdf":
        content = get_pdf_text_from_bytes(data)
        doc_type = "pdf"
    elif extension in (".docx", ".doc"):
        content = get_docx_text_from_bytes(data)
        doc_type = "docx"
    else:
        raise ValueError(f"Desteklenmeyen dosya formati: {extension}")

    if not content:
        raise ValueError("Dokumandan metin cikarilamadi")

    return {
        "filename": filename,
        "content": content,
        "doc_type": doc_type,
        "checksum": checksum_bytes(data),
    }
